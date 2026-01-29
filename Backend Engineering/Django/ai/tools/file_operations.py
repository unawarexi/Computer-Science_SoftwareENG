"""
File Operations Tool
====================
Safe file operations for AI agents.

Provides sandboxed file access with:
- Directory restrictions
- Read/write permissions
- Size limits
- Path validation
"""

import logging
import os
from pathlib import Path
from typing import Any, Dict, List, Optional

from langchain_core.tools import tool

from ai.tools.base import BaseTool

logger = logging.getLogger("ai.tools.files")


# =============================================================================
# FILE OPERATIONS TOOL
# =============================================================================

class FileOperationsTool(BaseTool):
    """
    Safe file operations within allowed directories.
    
    Features:
    - Sandboxed to specific directories
    - Read, write, list, delete operations
    - File size limits
    - Path traversal protection
    """
    
    def __init__(
        self,
        allowed_directories: List[str],
        max_file_size: int = 10_000_000,  # 10MB
        read_only: bool = False,
    ):
        """
        Initialize file operations tool.
        
        Args:
            allowed_directories: List of directories that can be accessed
            max_file_size: Maximum file size in bytes
            read_only: If True, only allow read operations
        """
        super().__init__()
        self.allowed_directories = [
            Path(d).resolve() for d in allowed_directories
        ]
        self.max_file_size = max_file_size
        self.read_only = read_only
    
    @property
    def name(self) -> str:
        return "file_operations"
    
    @property
    def description(self) -> str:
        return """Perform file operations (read, write, list, delete).
        Operations are restricted to allowed directories for security."""
    
    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "operation": {
                "type": "string",
                "description": "Operation: 'read', 'write', 'list', 'delete', 'exists'",
            },
            "path": {
                "type": "string",
                "description": "File or directory path",
            },
            "content": {
                "type": "string",
                "description": "Content to write (for write operation)",
            },
        }
    
    @property
    def required_params(self) -> List[str]:
        return ["operation", "path"]
    
    def _validate_path(self, path: str) -> Path:
        """
        Validate and resolve path.
        
        Ensures path is within allowed directories.
        """
        resolved = Path(path).resolve()
        
        # Check if path is within allowed directories
        for allowed in self.allowed_directories:
            try:
                resolved.relative_to(allowed)
                return resolved
            except ValueError:
                continue
        
        raise PermissionError(
            f"Access denied: {path} is outside allowed directories. "
            f"Allowed: {[str(d) for d in self.allowed_directories]}"
        )
    
    async def execute(
        self,
        operation: str,
        path: str,
        content: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Execute file operation.
        
        Args:
            operation: Operation to perform
            path: File/directory path
            content: Content for write operations
            
        Returns:
            Operation result
        """
        try:
            validated_path = self._validate_path(path)
            
            if operation == "read":
                return await self._read_file(validated_path)
            elif operation == "write":
                if self.read_only:
                    raise PermissionError("Write operations not allowed")
                return await self._write_file(validated_path, content or "")
            elif operation == "list":
                return await self._list_directory(validated_path)
            elif operation == "delete":
                if self.read_only:
                    raise PermissionError("Delete operations not allowed")
                return await self._delete_file(validated_path)
            elif operation == "exists":
                return {"exists": validated_path.exists()}
            else:
                return {"error": f"Unknown operation: {operation}"}
                
        except PermissionError as e:
            return {"error": str(e)}
        except Exception as e:
            logger.error(f"File operation failed: {e}")
            return {"error": str(e)}
    
    async def _read_file(self, path: Path) -> Dict[str, Any]:
        """Read file contents."""
        if not path.exists():
            return {"error": f"File not found: {path}"}
        
        if not path.is_file():
            return {"error": f"Not a file: {path}"}
        
        # Check file size
        size = path.stat().st_size
        if size > self.max_file_size:
            return {
                "error": f"File too large: {size} bytes (max: {self.max_file_size})"
            }
        
        try:
            content = path.read_text()
            return {
                "content": content,
                "size": size,
                "path": str(path),
            }
        except UnicodeDecodeError:
            # Try binary read
            content = path.read_bytes()
            return {
                "content": f"<binary file, {len(content)} bytes>",
                "size": size,
                "path": str(path),
                "binary": True,
            }
    
    async def _write_file(self, path: Path, content: str) -> Dict[str, Any]:
        """Write content to file."""
        # Check content size
        if len(content.encode()) > self.max_file_size:
            return {"error": f"Content too large (max: {self.max_file_size} bytes)"}
        
        # Create parent directories
        path.parent.mkdir(parents=True, exist_ok=True)
        
        # Write file
        path.write_text(content)
        
        return {
            "success": True,
            "path": str(path),
            "size": len(content),
        }
    
    async def _list_directory(self, path: Path) -> Dict[str, Any]:
        """List directory contents."""
        if not path.exists():
            return {"error": f"Directory not found: {path}"}
        
        if not path.is_dir():
            return {"error": f"Not a directory: {path}"}
        
        items = []
        for item in path.iterdir():
            try:
                stat = item.stat()
                items.append({
                    "name": item.name,
                    "type": "directory" if item.is_dir() else "file",
                    "size": stat.st_size if item.is_file() else None,
                })
            except Exception:
                items.append({
                    "name": item.name,
                    "type": "unknown",
                })
        
        return {
            "path": str(path),
            "items": items,
            "count": len(items),
        }
    
    async def _delete_file(self, path: Path) -> Dict[str, Any]:
        """Delete file."""
        if not path.exists():
            return {"error": f"File not found: {path}"}
        
        if path.is_dir():
            return {"error": "Cannot delete directories"}
        
        path.unlink()
        
        return {
            "success": True,
            "deleted": str(path),
        }


# =============================================================================
# DOCUMENT LOADER TOOL
# =============================================================================

class DocumentLoaderTool(BaseTool):
    """
    Load and parse various document formats.
    
    Supports: TXT, PDF, DOCX, Markdown, JSON, CSV
    """
    
    SUPPORTED_FORMATS = {
        ".txt": "text",
        ".md": "markdown",
        ".json": "json",
        ".csv": "csv",
        ".pdf": "pdf",
        ".docx": "docx",
    }
    
    def __init__(
        self,
        allowed_directories: List[str],
        max_file_size: int = 50_000_000,  # 50MB
    ):
        """
        Initialize document loader.
        
        Args:
            allowed_directories: Directories to access
            max_file_size: Maximum file size
        """
        super().__init__()
        self.file_ops = FileOperationsTool(
            allowed_directories=allowed_directories,
            max_file_size=max_file_size,
            read_only=True,
        )
    
    @property
    def name(self) -> str:
        return "document_loader"
    
    @property
    def description(self) -> str:
        return """Load and parse documents (TXT, MD, JSON, CSV, PDF, DOCX).
        Returns document content in a structured format."""
    
    @property
    def parameters(self) -> Dict[str, Any]:
        return {
            "path": {
                "type": "string",
                "description": "Path to the document",
            },
        }
    
    @property
    def required_params(self) -> List[str]:
        return ["path"]
    
    async def execute(self, path: str) -> Dict[str, Any]:
        """
        Load and parse a document.
        
        Args:
            path: Document path
            
        Returns:
            Parsed document content
        """
        file_path = Path(path)
        extension = file_path.suffix.lower()
        
        if extension not in self.SUPPORTED_FORMATS:
            return {
                "error": f"Unsupported format: {extension}. "
                         f"Supported: {list(self.SUPPORTED_FORMATS.keys())}"
            }
        
        format_type = self.SUPPORTED_FORMATS[extension]
        
        if format_type in ("text", "markdown"):
            return await self._load_text(path)
        elif format_type == "json":
            return await self._load_json(path)
        elif format_type == "csv":
            return await self._load_csv(path)
        elif format_type == "pdf":
            return await self._load_pdf(path)
        elif format_type == "docx":
            return await self._load_docx(path)
        
        return {"error": "Format not implemented"}
    
    async def _load_text(self, path: str) -> Dict[str, Any]:
        """Load text file."""
        result = await self.file_ops.execute("read", path)
        if "error" in result:
            return result
        
        return {
            "format": "text",
            "content": result["content"],
            "size": result["size"],
        }
    
    async def _load_json(self, path: str) -> Dict[str, Any]:
        """Load JSON file."""
        import json
        
        result = await self.file_ops.execute("read", path)
        if "error" in result:
            return result
        
        try:
            data = json.loads(result["content"])
            return {
                "format": "json",
                "content": data,
                "size": result["size"],
            }
        except json.JSONDecodeError as e:
            return {"error": f"Invalid JSON: {e}"}
    
    async def _load_csv(self, path: str) -> Dict[str, Any]:
        """Load CSV file."""
        import csv
        import io
        
        result = await self.file_ops.execute("read", path)
        if "error" in result:
            return result
        
        reader = csv.DictReader(io.StringIO(result["content"]))
        rows = list(reader)
        
        return {
            "format": "csv",
            "headers": reader.fieldnames,
            "rows": rows,
            "row_count": len(rows),
        }
    
    async def _load_pdf(self, path: str) -> Dict[str, Any]:
        """Load PDF file."""
        try:
            import pypdf
            
            reader = pypdf.PdfReader(path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "format": "pdf",
                "content": text,
                "pages": len(reader.pages),
            }
        except ImportError:
            return {"error": "pypdf not installed"}
        except Exception as e:
            return {"error": f"PDF parsing failed: {e}"}
    
    async def _load_docx(self, path: str) -> Dict[str, Any]:
        """Load DOCX file."""
        try:
            from docx import Document
            
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)
            
            return {
                "format": "docx",
                "content": text,
                "paragraphs": len(doc.paragraphs),
            }
        except ImportError:
            return {"error": "python-docx not installed"}
        except Exception as e:
            return {"error": f"DOCX parsing failed: {e}"}


# =============================================================================
# LANGCHAIN TOOL DECORATORS
# =============================================================================

@tool
def read_file_content(path: str) -> str:
    """Read the contents of a text file."""
    import asyncio
    from django.conf import settings
    
    # Use media root as allowed directory
    allowed_dirs = [settings.MEDIA_ROOT] if hasattr(settings, 'MEDIA_ROOT') else ["/tmp"]
    tool = FileOperationsTool(allowed_directories=allowed_dirs, read_only=True)
    result = asyncio.get_event_loop().run_until_complete(
        tool.execute(operation="read", path=path)
    )
    
    if "error" in result:
        return f"Error: {result['error']}"
    return result.get("content", "")


@tool
def list_directory(path: str) -> str:
    """List files and folders in a directory."""
    import asyncio
    import json
    from django.conf import settings
    
    allowed_dirs = [settings.MEDIA_ROOT] if hasattr(settings, 'MEDIA_ROOT') else ["/tmp"]
    tool = FileOperationsTool(allowed_directories=allowed_dirs, read_only=True)
    result = asyncio.get_event_loop().run_until_complete(
        tool.execute(operation="list", path=path)
    )
    
    if "error" in result:
        return f"Error: {result['error']}"
    
    items = result.get("items", [])
    return json.dumps(items, indent=2)
